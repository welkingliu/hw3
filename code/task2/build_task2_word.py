from __future__ import annotations

import csv
import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/welkinliu/Desktop/hw3")
RUNS = ROOT / "code" / "runs"
OUT = ROOT / "outputs" / "task2_report"
DOCX = OUT / "任务二_ACT跨环境泛化实验报告.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) <= 10 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9.5)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, "F2F4F7")
        set_text(cell, h, bold=True)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_text(cells[i], val)
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if widths:
        for row in table.rows:
            for cell, width in zip(row.cells, widths):
                cell.width = Inches(width)
    doc.add_paragraph()


def add_figure(doc: Document, path: Path, caption: str, width: float = 5.8) -> None:
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(90, 90, 90)


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "E8EEF5")
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor(31, 77, 120)
    p.add_run("\n" + body)


def read_clean_metrics(name: str) -> list[dict[str, float]]:
    rows_by_epoch: dict[int, dict[str, float]] = {}
    with (RUNS / name / "metrics.csv").open(newline="", encoding="utf-8") as f:
        for row in csv.DictReader(f):
            epoch = int(row["epoch"])
            rows_by_epoch[epoch] = {
                "epoch": epoch,
                "train_l1": float(row["train_l1"]),
                "val_l1": float(row["val_l1"]),
            }
    return [rows_by_epoch[k] for k in sorted(rows_by_epoch)]


def style_doc(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.12
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color in [
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(31, 77, 120)),
        ("Heading 3", 11.5, RGBColor(31, 77, 120)),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("任务二实验报告")
    r.bold = True
    r.font.size = Pt(22)
    r.font.name = "Calibri"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.color.rgb = RGBColor(31, 77, 120)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("基于 LeRobot / ACT 的跨环境泛化实验")
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(80, 80, 80)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("结果目录：/Users/welkinliu/Desktop/hw3/code/runs")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(100, 100, 100)


def pct_improve(base: float, new: float) -> float:
    return (base - new) / base * 100.0


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    summary = json.loads((OUT / "task2_summary.json").read_text())
    cfg = json.loads((RUNS / "act_A" / "config.json").read_text())
    a_final = summary["act_A_final"]
    abc_final = summary["act_ABC_final"]
    eval_a = summary["eval_A_to_D"]
    eval_abc = summary["eval_ABC_to_D"]

    doc = Document()
    style_doc(doc)
    add_title(doc)

    add_callout(
        doc,
        "摘要",
        "本报告对应任务二：先训练只使用环境 A 的 ACT 基础模型，再训练混合环境 A/B/C 的 ACT 多环境模型，最后将两个模型在未见过的环境 D 上进行 zero-shot 测试。实验结果显示，多环境训练模型在验证集和环境 D 测试中的 Action L1 均更低，体现了更好的跨环境泛化能力。",
    )

    doc.add_heading("1. 作业要求与满足逻辑", level=1)
    add_table(
        doc,
        ["要求", "本项目实现", "结果材料"],
        [
            ["环境 A 基础训练", "使用环境 A 数据训练 ACT-A 模型。", "runs/act_A/best.pt"],
            ["环境 A/B/C 联合训练", "混合 A、B、C 三个环境数据，使用相同网络和超参数训练 ACT-ABC。", "runs/act_ABC/best.pt"],
            ["环境 D zero-shot 测试", "两个模型均不微调，直接在环境 D 上测试。", "runs/eval_*_to_D/results.json"],
            ["指标对比", "统计 Action L1 Loss 和 Success Rate，并绘制训练曲线。", "loss_curve.png / task2_*.png"],
            ["鲁棒性分析", "比较单环境训练与多环境训练在视觉分布偏移下的表现。", "报告第 5 节"],
        ],
        widths=[1.55, 3.55, 1.45],
    )

    doc.add_heading("2. 实现方法", level=1)
    doc.add_paragraph("当前代码位于 code/ 目录。由于本地 data/ 没有真实 CALVIN 数据，本项目实现了一个轻量 CALVIN-like 多环境数据集，用不同背景颜色、物体颜色、噪声和风格偏移模拟 A/B/C/D 环境中的视觉分布差异。模型结构采用 ACT-style action chunking 策略网络：CNN 视觉编码器 + MLP 状态编码器 + Transformer Decoder，并一次性预测长度为 8 的动作序列。")
    add_table(
        doc,
        ["模块", "路径", "作用"],
        [
            ["数据集", "src/act_hw3/synthetic_calvin.py", "生成 A/B/C/D 多环境视觉-动作样本"],
            ["模型", "src/act_hw3/model.py", "ACT-style 策略网络，预测 action chunk"],
            ["训练", "src/act_hw3/train.py", "训练 ACT-A 与 ACT-ABC"],
            ["测试", "src/act_hw3/evaluate.py", "在环境 D 上进行 zero-shot 评估"],
            ["一键实验", "run_experiment.py", "串联训练与测试流程"],
        ],
        widths=[0.95, 2.15, 3.25],
    )

    doc.add_heading("3. 实验设置", level=1)
    add_table(
        doc,
        ["项目", "设置"],
        [
            ["Network Architecture", "CNN image encoder + MLP state encoder + Transformer Decoder"],
            ["Image Size", f"{cfg['image_size']} x {cfg['image_size']}"],
            ["State Dim / Action Dim", f"{cfg['state_dim']} / {cfg['action_dim']}"],
            ["Action Chunk Size", str(cfg["chunk_size"])],
            ["Batch Size", str(cfg["batch_size"])],
            ["Learning Rate", str(cfg["learning_rate"])],
            ["Optimizer", "AdamW"],
            ["Epochs", str(cfg["epochs"])],
            ["Loss Function", "Action L1 Loss"],
            ["Success Threshold", str(cfg["success_threshold"])],
            ["Device", cfg["device"]],
        ],
        widths=[2.1, 4.1],
    )

    doc.add_heading("4. 训练结果", level=1)
    add_table(
        doc,
        ["模型", "训练环境", "Final Train L1", "Final Val L1", "权重"],
        [
            ["ACT-A", "A", f"{a_final['train_l1']:.4f}", f"{a_final['val_l1']:.4f}", "runs/act_A/best.pt"],
            ["ACT-ABC", "A, B, C", f"{abc_final['train_l1']:.4f}", f"{abc_final['val_l1']:.4f}", "runs/act_ABC/best.pt"],
        ],
        widths=[0.9, 1.1, 1.2, 1.2, 2.0],
    )
    doc.add_paragraph(f"从验证集 Action L1 看，ACT-ABC 的最终 Val L1 为 {abc_final['val_l1']:.4f}，低于 ACT-A 的 {a_final['val_l1']:.4f}，相对降低约 {pct_improve(a_final['val_l1'], abc_final['val_l1']):.1f}%。")
    add_figure(doc, OUT / "task2_val_loss_compare.png", "图 1：ACT-A 与 ACT-ABC 验证集 Action L1 曲线")
    add_figure(doc, RUNS / "act_A" / "loss_curve.png", "图 2：ACT-A 训练曲线", width=5.4)
    add_figure(doc, RUNS / "act_ABC" / "loss_curve.png", "图 3：ACT-ABC 训练曲线", width=5.4)

    doc.add_heading("5. 环境 D Zero-shot 测试", level=1)
    add_table(
        doc,
        ["模型", "训练环境", "测试环境", "Action L1", "Success Rate", "测试样本数"],
        [
            ["ACT-A", "A", "D", f"{eval_a['action_l1']:.4f}", f"{eval_a['success_rate']:.3f}", str(eval_a["num_samples"])],
            ["ACT-ABC", "A, B, C", "D", f"{eval_abc['action_l1']:.4f}", f"{eval_abc['success_rate']:.3f}", str(eval_abc["num_samples"])],
        ],
        widths=[0.8, 1.0, 0.8, 1.0, 1.1, 1.0],
    )
    doc.add_paragraph(f"在未见过的环境 D 上，ACT-ABC 的 Action L1 为 {eval_abc['action_l1']:.4f}，低于 ACT-A 的 {eval_a['action_l1']:.4f}，相对降低约 {pct_improve(eval_a['action_l1'], eval_abc['action_l1']):.1f}%。两个模型的 Success Rate 均为 1.0，说明在当前阈值 0.16 下均能完成任务，但多环境模型的动作误差更小。")
    add_figure(doc, OUT / "task2_zero_shot_l1.png", "图 4：环境 D zero-shot Action L1 对比", width=5.0)

    doc.add_heading("6. 分析与结论", level=1)
    doc.add_paragraph("ACT 的核心是 action chunking：模型不是只预测单步动作，而是一次预测一段动作序列。这样可以降低逐步控制中的短期抖动，并让策略在视觉输入存在一定偏移时仍能保持动作连续性。")
    doc.add_paragraph("从实验结果看，ACT-ABC 同时接触 A/B/C 三种视觉风格，因此在 D 环境上具有更低的 Action L1。这说明多环境训练能提升模型对视觉分布偏移的适应能力。单环境 ACT-A 只学习环境 A 的视觉模式，虽然在当前合成任务中也能达到较高成功率，但预测动作误差更大。")

    doc.add_heading("7. 局限与改进点", level=1)
    add_table(
        doc,
        ["问题", "说明", "改进方向"],
        [
            ["数据集为合成 CALVIN-like", "当前没有接入真实 CALVIN/LeRobot 数据。", "替换为真实 CALVIN 数据加载器。"],
            ["Success Rate 是阈值近似", "当前用 Action L1 是否低于阈值近似成功率。", "接入真实仿真 rollout 评估成功率。"],
            ["环境偏移较可控", "A/B/C/D 的偏移由合成样式控制。", "增加真实视觉扰动、遮挡、光照变化。"],
            ["未使用官方 LeRobot ACT", "当前为轻量 ACT-style 实现。", "迁移到 LeRobot 官方 ACT 训练脚本。"],
        ],
        widths=[1.4, 2.35, 2.65],
    )

    doc.add_heading("8. 可提交材料清单", level=1)
    add_table(
        doc,
        ["材料", "路径"],
        [
            ["ACT-A 权重", "code/runs/act_A/best.pt"],
            ["ACT-ABC 权重", "code/runs/act_ABC/best.pt"],
            ["ACT-A 训练曲线", "code/runs/act_A/loss_curve.png"],
            ["ACT-ABC 训练曲线", "code/runs/act_ABC/loss_curve.png"],
            ["环境 D 测试结果", "code/runs/eval_A_to_D/results.json / eval_ABC_to_D/results.json"],
            ["训练与测试代码", "code/src/act_hw3/"],
        ],
        widths=[2.0, 4.4],
    )

    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    main()
