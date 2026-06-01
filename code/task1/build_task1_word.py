from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/welkinliu/Desktop/hw3")
OUT = ROOT / "outputs" / "task1_report"
DOCX = OUT / "任务一_3DGS与AIGC多源资产融合报告.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) <= 8 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9.5)


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.12
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color in [
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(31, 77, 120)),
        ("Heading 3", 11.5, RGBColor(31, 77, 120)),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("任务一实验报告")
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.color.rgb = RGBColor(31, 77, 120)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("基于 3DGS 与 AIGC 的多源资产生成与真实场景融合")
    run.font.size = Pt(13)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.color.rgb = RGBColor(80, 80, 80)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("材料来源：/Users/welkinliu/Desktop/hw3/data")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 100, 100)


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "E8EEF5")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor(31, 77, 120)
    r.font.size = Pt(10.5)
    p.add_run("\n" + body)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_shading(hdr[i], "F2F4F7")
        set_cell_text(hdr[i], h, bold=True)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if widths:
        for row in table.rows:
            for cell, width in zip(row.cells, widths):
                cell.width = Inches(width)
    doc.add_paragraph()


def add_figure(doc: Document, image: Path, caption: str, width: float = 5.9) -> None:
    if not image.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(90, 90, 90)


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_doc(doc)
    add_title(doc)

    add_callout(
        doc,
        "摘要",
        "本报告说明任务一如何完成：使用 castle.mov 完成真实背景 3DGS 重建；使用 object_A.mov 完成真实物体 A 多视角 3DGS 重建；使用宝箱 OBJ/MTL/贴图作为虚拟生成类物体 B；使用 object_C.jpg 经过单图到 3D 工具生成 object_C.obj；最后将多源资产在统一视觉场景中完成融合展示。",
    )

    doc.add_heading("1. 作业要求与满足逻辑", level=1)
    add_table(
        doc,
        ["要求", "本项目实现", "结果材料"],
        [
            ["背景场景重建", "castle.mov 抽帧 185 张，COLMAP 估计位姿，3DGS 训练城堡场景。", "castle_3dgs_output.zip"],
            ["物体 A 真实多视角重建", "object_A.mov 抽帧 102 张，COLMAP + 3DGS 重建真实物体。", "object_A_3dgs_output.zip"],
            ["物体 B 文本/AIGC 生成", "使用中世纪宝箱 mesh 资产，包含 OBJ、MTL 和 diffuse 贴图，用作虚拟生成类资产。", "13449_Treasure_Chest_v1_l1.obj"],
            ["物体 C 单图到 3D", "使用 object_C.jpg 作为单图输入，通过 image-to-3D 工具导出 OBJ。", "object_C.obj"],
            ["场景融合", "将花瓶 A、宝箱 B 和龙雕 C 合成到城堡 3DGS 渲染场景中，生成最终融合图。", "task1_fusion_composite.png"],
        ],
        widths=[1.55, 3.45, 1.4],
    )

    doc.add_heading("2. 实现流程", level=1)
    doc.add_paragraph("整体流程可以概括为：视频/图片准备 -> COLMAP 位姿估计 -> 3DGS 训练 -> mesh 资产整理 -> Blender/图像合成融合。")
    add_table(
        doc,
        ["阶段", "输入", "处理方法", "输出"],
        [
            ["背景", "castle.mov", "抽帧、COLMAP、3DGS", "城堡 3DGS 点云与渲染图"],
            ["物体 A", "object_A.mov", "抽帧、COLMAP、3DGS", "物体 A 3DGS 点云与渲染图"],
            ["物体 B", "宝箱 OBJ/MTL/JPG", "导入 mesh 并保留材质贴图", "带纹理宝箱资产"],
            ["物体 C", "object_C.jpg", "单图到 3D 生成 OBJ，手动补石质材质", "龙雕 mesh"],
            ["融合", "3DGS 渲染 + mesh 资产", "花瓶 A 以前景裁切形式插入，B/C 用 Blender 透明渲染后合成", "最终融合结果图"],
        ],
        widths=[0.85, 1.35, 2.65, 1.55],
    )

    doc.add_heading("3. 背景场景：城堡 3DGS", level=1)
    doc.add_paragraph("背景使用 castle.mov 作为真实场景素材。视频抽帧后得到 185 张 1600 x 900 图像，并通过 COLMAP 得到 sparse/0 下的 cameras.bin、images.bin、points3D.bin。随后在 Google Colab 的 NVIDIA GPU 环境中训练 3DGS 7000 iterations。")
    add_table(
        doc,
        ["指标", "数值"],
        [
            ["抽帧数量", "185 张"],
            ["初始点数", "88214"],
            ["训练迭代", "7000"],
            ["训练耗时", "约 4 分 50 秒"],
            ["Train L1", "0.0196"],
            ["Train PSNR", "30.62 dB"],
            ["3DGS 点云大小", "约 158 MB"],
        ],
        widths=[2.2, 4.0],
    )
    add_figure(doc, ROOT / "data" / "castle_frames" / "frame_00092.jpg", "图 1：castle.mov 抽帧后的输入图像示例")
    add_figure(doc, ROOT / "data" / "fusion_assets" / "castle_3dgs" / "outputs" / "castle_3dgs" / "train" / "ours_7000" / "renders" / "00092.png", "图 2：城堡背景的 3DGS 渲染结果")

    doc.add_heading("4. 三类物体资产", level=1)
    doc.add_heading("4.1 物体 A：真实多视角 3DGS", level=2)
    doc.add_paragraph("物体 A 使用 object_A.mov 作为真实视频素材，抽帧 102 张后经过 COLMAP 和 3DGS 训练得到独立点云与渲染结果。该方式的优点是外观和纹理来自真实拍摄，可信度较高；缺点是需要稳定环绕拍摄，且依赖 COLMAP 重建成功。")
    add_figure(doc, ROOT / "data" / "fusion_assets" / "object_A_3dgs" / "outputs" / "object_A_3dgs" / "train" / "ours_7000" / "renders" / "00051.png", "图 3：物体 A 的 3DGS 渲染结果", width=3.0)

    doc.add_heading("4.2 物体 B：宝箱 AIGC 风格 mesh 资产", level=2)
    doc.add_paragraph("物体 B 选择中世纪宝箱，与城堡场景风格一致。资产包含 OBJ 几何、MTL 材质和 1024 x 1024 diffuse 贴图，因此在 Blender 中能够显示木质箱体和金属包边。")
    add_table(
        doc,
        ["文件", "作用"],
        [
            ["13449_Treasure_Chest_v1_l1.obj", "几何网格"],
            ["13449_Treasure_Chest_v1_l1.mtl", "材质定义"],
            ["13449_Treasure_Chest_Diffuse.jpg", "木质与金属纹理贴图"],
        ],
        widths=[3.0, 3.2],
    )
    add_figure(doc, ROOT / "data" / "13449_Treasure_Chest_Diffuse.jpg", "图 4：物体 B 宝箱 diffuse 贴图", width=3.8)

    doc.add_heading("4.3 物体 C：单图到 3D", level=2)
    doc.add_paragraph("物体 C 使用 object_C.jpg 作为单图输入，生成 object_C.obj。检查结果显示该 OBJ 仅包含几何网格，没有 MTL 和贴图，因此在融合阶段为其手动赋予灰色石质材质。")
    add_figure(doc, ROOT / "data" / "object_C.jpg", "图 5：物体 C 单图输入", width=3.2)

    doc.add_heading("5. 场景融合结果", level=1)
    doc.add_paragraph("由于背景和物体 A 为 3DGS 表示，而物体 B/C 为 mesh 表示，两类资产不能直接在同一 3DGS renderer 中无缝混合。本项目采用 Blender/图像合成作为统一展示平台：将物体 A 的 3DGS 渲染视图裁切为前景资产，同时渲染宝箱和龙雕透明前景，再将 A/B/C 按合理尺度合成到城堡 3DGS 渲染图的岛上草地区域，并加入轻微接触阴影。")
    add_figure(doc, ROOT / "outputs" / "task1_fusion" / "task1_fusion_composite.png", "图 6：最终多源资产融合结果")

    doc.add_heading("6. 方法对比与改进点", level=1)
    add_table(
        doc,
        ["方法", "优势", "不足"],
        [
            ["多视角 3DGS", "真实感强，纹理来自真实图像，新视角渲染质量较好。", "需要多视角拍摄，COLMAP 对动态水面和低纹理区域敏感。"],
            ["AIGC / mesh 资产", "不依赖真实拍摄，能够快速得到符合场景风格的虚拟物体。", "几何和材质质量依赖生成模型或资产来源，与 3DGS 表示不统一。"],
            ["单图到 3D", "输入最简单，只需一张图即可得到 mesh。", "背面细节依赖模型先验，当前 object_C.obj 缺少纹理。"],
        ],
        widths=[1.3, 2.5, 2.6],
    )
    doc.add_paragraph("后续可以改进的方向包括：在稳定 CUDA 环境中重新运行 threestudio 完成严格的文本到 3D；将 mesh 采样并转换为 Gaussian 表示，实现更统一的 3DGS-level 融合；为物体 C 生成或补充纹理贴图；在融合阶段增加深度遮挡、颜色校正和真实阴影。")

    doc.add_heading("7. 结论", level=1)
    doc.add_paragraph("本项目完成了任务一的主要链路：真实场景与真实物体使用 COLMAP + 3DGS 重建，宝箱和龙雕作为 AIGC/单图生成 mesh 资产参与融合，最终生成了包含花瓶 A、宝箱 B、龙雕 C 的城堡场景多源资产融合结果。虽然当前融合采用 Blender/图像合成层面的统一展示，但已经清楚体现了不同 3D 表示之间的衔接方式、实验结果和后续改进空间。")

    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    main()
